# -*- coding: utf-8 -*-
"""
classify_weekly_onefile.py

Weekly pipeline:
1) Pick latest weekly CSV in output/weekly (or -i).
2) (Optional) Google Translate title/abstract into title_zh/abstract_zh unless --skip-translate.
3) Hybrid classify:
   - Parse category descriptions from classification.txt (format: 类别名：描述)
   - Extract include/exclude keyword hints heuristically from “包括/不包括” segments
   - Rule layer:
       * If exclude keywords hit -> that category is blocked
       * If include keywords hit strongly -> auto-assign (no online LLM)
   - Transformer semantic routing (sentence-transformers) to get topK candidates
   - Gemini final confirmation using by_id mapping output; retry missing ids once
   - Second-missing ids go to a dedicated Excel sheet
4) Output:
   - output/weekly/<input_stem><output_suffix>.xlsx (each category per sheet; items may repeat)
   - output/weekly/<input_stem><output_suffix>.docx (grouped by category blocks)

Environment variables:
  GOOGLE_AI_API_KEY or GEMINI_API_KEY (required)
  GOOGLE_AI_MODEL or GEMINI_MODEL (default gemini-2.5-flash)
  GOOGLE_AI_TEMPERATURE (default 0)
  GOOGLE_AI_MAX_RETRIES (default 3)
  GOOGLE_AI_RETRY_BASE_SECONDS (default 3)
  GOOGLE_AI_RETRY_MAX_SECONDS (default 60)
  GOOGLE_AI_BASE_URL (optional)

  # Translation
  TRANSLATE_PROVIDER (default google; choices: google, none)
  MAX_ABSTRACT_CHARS_TO_TRANSLATE (default 1600)
  TRANSLATE_BATCH_SIZE_TITLE (default 12)
  TRANSLATE_BATCH_SIZE_ABSTRACT (default 3)

  # Classification
  CLASSIFY_BATCH_SIZE (default 12)
  KEYWORD_STRONG_HITS (default 2)   # include hits >= this -> auto-assign
  TOPK_CANDIDATES (default 4)
  EMB_STRONG_THRESHOLD (default 0.60)  # if top1 >= this -> auto-assign (if not blocked)
  EMB_WEAK_THRESHOLD (default 0.45)    # candidates must be >= this

  # Embedding model
  EMB_MODEL_NAME (default "sentence-transformers/all-MiniLM-L6-v2")

Notes:
- classification.txt is optimized for an LLM; embeddings do NOT reliably “understand” negation,
  so we strip "不包括..." part for embedding prototypes.
"""

import argparse
import hashlib
import json
import os
import random
import re
import time
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


print("[boot] classify_weekly_onefile.py started", flush=True)

# ============================
# Google AI Studio / Gemini configuration
# ============================
GOOGLE_AI_MODEL = (
    os.getenv("GOOGLE_AI_MODEL")
    or os.getenv("GEMINI_MODEL")
    or "gemini-2.5-flash"
).strip()
GOOGLE_AI_TEMPERATURE = float(os.getenv("GOOGLE_AI_TEMPERATURE", os.getenv("GEMINI_TEMPERATURE", "0")))
GOOGLE_AI_MAX_RETRIES = int(os.getenv("GOOGLE_AI_MAX_RETRIES", os.getenv("GEMINI_MAX_RETRIES", "3")))
GOOGLE_AI_RETRY_BASE_SECONDS = float(os.getenv("GOOGLE_AI_RETRY_BASE_SECONDS", os.getenv("GEMINI_RETRY_BASE_SECONDS", "3")))
GOOGLE_AI_RETRY_MAX_SECONDS = float(os.getenv("GOOGLE_AI_RETRY_MAX_SECONDS", os.getenv("GEMINI_RETRY_MAX_SECONDS", "60")))
GOOGLE_AI_BASE_URL = (
    os.getenv("GOOGLE_AI_BASE_URL")
    or os.getenv("GEMINI_BASE_URL")
    or "https://generativelanguage.googleapis.com/v1beta"
).strip().rstrip("/")
GOOGLE_AI_TIMEOUT = int(os.getenv("GOOGLE_AI_TIMEOUT", os.getenv("GEMINI_TIMEOUT", "120")))
GEMINI_TRANSLATE_MAX_RETRIES = int(os.getenv("GEMINI_TRANSLATE_MAX_RETRIES", "2"))
GEMINI_CLASSIFY_MAX_RETRIES = int(os.getenv("GEMINI_CLASSIFY_MAX_RETRIES", os.getenv("GOOGLE_AI_MAX_RETRIES", "3")))

MAX_ABSTRACT_CHARS_TO_TRANSLATE = int(os.getenv("MAX_ABSTRACT_CHARS_TO_TRANSLATE", "1600"))
TRANSLATE_BATCH_SIZE_TITLE = int(os.getenv("TRANSLATE_BATCH_SIZE_TITLE", "12"))
TRANSLATE_BATCH_SIZE_ABSTRACT = int(os.getenv("TRANSLATE_BATCH_SIZE_ABSTRACT", "3"))
TRANSLATE_PROVIDER = os.getenv("TRANSLATE_PROVIDER", "google").strip().lower()
GEMINI_TRANSLATE_FALLBACK_PROVIDER = os.getenv("GEMINI_TRANSLATE_FALLBACK_PROVIDER", "google").strip().lower()
GOOGLE_TRANSLATE_MAX_RETRIES = int(os.getenv("GOOGLE_TRANSLATE_MAX_RETRIES", "4"))

CLASSIFY_BATCH_SIZE = int(os.getenv("CLASSIFY_BATCH_SIZE", "12"))
KEYWORD_STRONG_HITS = int(os.getenv("KEYWORD_STRONG_HITS", "2"))
TOPK_CANDIDATES = int(os.getenv("TOPK_CANDIDATES", "4"))
EMB_STRONG_THRESHOLD = float(os.getenv("EMB_STRONG_THRESHOLD", "0.60"))
EMB_WEAK_THRESHOLD = float(os.getenv("EMB_WEAK_THRESHOLD", "0.45"))

EMB_MODEL_NAME = os.getenv("EMB_MODEL_NAME", "sentence-transformers/all-MiniLM-L6-v2").strip()

DEFAULT_CLASSIFICATION_FILE = "classification.txt"


# ============================
# Utils
# ============================
def _fmt_secs(s: float) -> str:
    s = max(0, int(s))
    h = s // 3600
    m = (s % 3600) // 60
    sec = s % 60
    if h > 0:
        return f"{h}h {m}m {sec}s"
    if m > 0:
        return f"{m}m {sec}s"
    return f"{sec}s"


def _chunked(items, n):
    for i in range(0, len(items), n):
        yield items[i: i + n]


def _normalize_cell(v) -> str:
    if v is None:
        return ""
    return str(v).strip()


def _clean_json_text(s: str) -> str:
    s = (s or "").strip()
    if s.startswith("```"):
        s = re.sub(r"^```(?:json)?\s*", "", s, flags=re.IGNORECASE)
        s = re.sub(r"\s*```$", "", s)
    return s.strip()


def _stable_id_from_row(row: Dict[str, Any]) -> str:
    """
    Prefer stable id from link. Fallback to title+source+published.
    """
    link = _normalize_cell(row.get("link", ""))
    title = _normalize_cell(row.get("title", ""))
    source = _normalize_cell(row.get("source", ""))
    pub = _normalize_cell(row.get("pub_date", "")) or _normalize_cell(row.get("published", ""))
    base = link if link else f"{title}||{source}||{pub}"
    h = hashlib.sha1(base.encode("utf-8", errors="ignore")).hexdigest()[:16]
    return h


def ensure_stable_id_column(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    records = df.to_dict(orient="records")
    df["stable_id"] = [_stable_id_from_row(r) for r in records]
    return df


def _atomic_write_text(path: Path, text: str):
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_name(path.name + ".tmp")
    tmp.write_text(text, encoding="utf-8")
    tmp.replace(path)


def _save_translation_checkpoint(df: pd.DataFrame, path: Path):
    cols = ["stable_id", "title_zh", "abstract_zh"]
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_name(path.name + ".tmp")
    df[cols].to_csv(tmp, index=False, encoding="utf-8-sig")
    tmp.replace(path)


def _load_translation_checkpoint(df: pd.DataFrame, path: Path) -> pd.DataFrame:
    if not path.exists():
        return df

    ckpt = pd.read_csv(path, encoding="utf-8-sig", keep_default_na=False)
    if "stable_id" not in ckpt.columns:
        return df

    applied = 0
    ckpt = ckpt.drop_duplicates("stable_id", keep="last").set_index("stable_id")
    for col in ["title_zh", "abstract_zh"]:
        if col not in ckpt.columns:
            continue
        for idx, sid in df["stable_id"].items():
            current = _normalize_cell(df.at[idx, col])
            cached = _normalize_cell(ckpt.at[sid, col]) if sid in ckpt.index else ""
            if not current and cached:
                df.at[idx, col] = cached
                applied += 1
    print(f"[checkpoint] loaded translation checkpoint: {path} (applied cells={applied})", flush=True)
    return df


def _load_classification_checkpoint(path: Path, allowed_labels: set) -> Dict[str, List[str]]:
    if not path.exists():
        return {}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception as e:
        print(f"[checkpoint] failed to load classification checkpoint {path}: {e}", flush=True)
        return {}

    raw = data.get("labels_by_id", data) if isinstance(data, dict) else {}
    out: Dict[str, List[str]] = {}
    if isinstance(raw, dict):
        for sid, labels in raw.items():
            if not isinstance(labels, list):
                continue
            clean = [_normalize_cell(x) for x in labels]
            clean = [x for x in clean if x in allowed_labels]
            out[str(sid)] = list(dict.fromkeys(clean))
    print(f"[checkpoint] loaded classification checkpoint: {path} (items={len(out)})", flush=True)
    return out


def _save_classification_checkpoint(path: Path, labels_by_id: Dict[str, List[str]]):
    payload = {
        "updated_at": datetime.utcnow().isoformat(timespec="seconds") + "Z",
        "labels_by_id": labels_by_id,
    }
    _atomic_write_text(path, json.dumps(payload, ensure_ascii=False, indent=2))


def _extract_date_from_name(filename: str):
    stem = Path(filename).stem
    m = re.search(r"(\d{4})[-_]?(\d{2})[-_]?(\d{2})", stem)
    if not m:
        return None
    y, mo, d = map(int, m.groups())
    try:
        return datetime(y, mo, d).date()
    except ValueError:
        return None


def pick_latest_weekly_csv(folder="output/weekly") -> Path:
    folder = Path(folder)
    if not folder.exists():
        raise FileNotFoundError(f"Folder not found: {folder.resolve()}")

    prefixes = ["weekly_news_with_abstract_", "news_with_abstract_"]
    candidates = []
    for p in folder.glob("*.csv"):
        if "_translated" in p.stem:
            continue
        if any(p.name.startswith(pref) for pref in prefixes):
            candidates.append(p)

    if not candidates:
        raise FileNotFoundError(f"No weekly CSV found in {folder.resolve()}")

    with_dates = []
    without_dates = []
    for p in candidates:
        d = _extract_date_from_name(p.name)
        if d is not None:
            with_dates.append((d, p))
        else:
            without_dates.append(p)

    if with_dates:
        with_dates.sort(key=lambda x: (x[0], x[1].stat().st_mtime), reverse=True)
        return with_dates[0][1]
    return max(without_dates, key=lambda p: p.stat().st_mtime)


def ensure_base_columns(df: pd.DataFrame) -> pd.DataFrame:
    needed = [
        "title",
        "link",
        "published",
        "source",
        "pub_date",
        "doi",
        "abstract",
        "abstract_source",
        "must_have_abstract",
        "title_zh",
        "abstract_zh",
        "categories",
    ]
    for c in needed:
        if c not in df.columns:
            df[c] = ""
        df[c] = df[c].fillna("").astype(str)
    return df


def build_text_for_classify(title: str, abstract: str, max_chars: int = 1600) -> str:
    """
    Keep it short-ish for embedding / LLM. Prefer abstract, fallback title.
    """
    title = _normalize_cell(title)
    abstract = _normalize_cell(abstract)
    if abstract:
        t = abstract
    else:
        t = title

    t = re.sub(r"\s+", " ", t).strip()
    if len(t) <= max_chars:
        return t
    # head + tail
    head = t[: max_chars // 2]
    tail = t[-max_chars // 2:]
    return head + " ... " + tail


# ============================
# Gemini client
# ============================
def _build_gemini_client():
    api_key = (
        os.getenv("GOOGLE_AI_API_KEY")
        or os.getenv("GEMINI_API_KEY")
        or ""
    ).strip()
    if not api_key:
        raise RuntimeError(
            "Missing GOOGLE_AI_API_KEY or GEMINI_API_KEY in environment. "
            "For GitHub Actions, add it in repo Settings -> Secrets and variables -> Actions."
        )

    if not GOOGLE_AI_MODEL:
        raise RuntimeError("GOOGLE_AI_MODEL is empty. Set GOOGLE_AI_MODEL or GEMINI_MODEL.")

    return {"api_key": api_key, "model": GOOGLE_AI_MODEL}


def _gemini_generate_content(client, messages, json_object: bool = True) -> str:
    import requests

    system_parts = []
    user_parts = []
    for msg in messages:
        role = msg.get("role", "user")
        content = msg.get("content", "")
        if role == "system":
            system_parts.append(str(content))
        else:
            user_parts.append(str(content))

    model = client["model"]
    model_path = model if model.startswith("models/") else f"models/{model}"
    url = f"{GOOGLE_AI_BASE_URL}/{model_path}:generateContent"

    payload = {
        "contents": [
            {
                "role": "user",
                "parts": [{"text": "\n\n".join(user_parts)}],
            }
        ],
        "generationConfig": {
            "temperature": GOOGLE_AI_TEMPERATURE,
        },
    }
    if system_parts:
        payload["systemInstruction"] = {"parts": [{"text": "\n\n".join(system_parts)}]}
    if json_object:
        payload["generationConfig"]["responseMimeType"] = "application/json"

    resp = requests.post(
        url,
        headers={
            "Content-Type": "application/json",
            "x-goog-api-key": client["api_key"],
        },
        json=payload,
        timeout=GOOGLE_AI_TIMEOUT,
    )
    try:
        resp.raise_for_status()
    except requests.HTTPError as e:
        raise RuntimeError(f"Gemini API HTTP {resp.status_code}: {resp.text[:1000]}") from e

    data = resp.json()
    candidates = data.get("candidates") or []
    if not candidates:
        raise RuntimeError(f"Gemini API returned no candidates: {json.dumps(data, ensure_ascii=False)[:1000]}")

    parts = candidates[0].get("content", {}).get("parts") or []
    text = "".join(str(p.get("text", "")) for p in parts).strip()
    if not text:
        raise RuntimeError(f"Gemini API returned empty content: {json.dumps(data, ensure_ascii=False)[:1000]}")
    return text


def _extract_content(resp) -> str:
    if isinstance(resp, str):
        return resp
    return resp.choices[0].message.content


def _call_with_retries(client, messages, label: str, json_object: bool = True, max_retries: Optional[int] = None):
    max_retries = max_retries or GOOGLE_AI_MAX_RETRIES
    last_err = None
    for attempt in range(1, max_retries + 1):
        try:
            return _gemini_generate_content(client, messages, json_object=json_object)
        except Exception as e:
            last_err = e
            wait = min(GOOGLE_AI_RETRY_MAX_SECONDS, GOOGLE_AI_RETRY_BASE_SECONDS * (2 ** (attempt - 1)))
            wait = wait + random.uniform(0, min(3.0, wait * 0.25))
            print(f"[gemini:{label}] error on attempt {attempt}: {e} (wait {wait:.1f}s)", flush=True)
            time.sleep(wait)
    raise RuntimeError(f"Gemini request failed after {max_retries} attempts: {last_err}")


# ============================
# Classification rules parsing
# ============================
@dataclass
class CategoryRule:
    name: str
    desc_llm: str          # full (with 不包括)
    desc_embed: str        # stripped (remove 不包括 segment)
    include_terms: List[str]  # extracted from 包括... segment (heuristic)
    exclude_terms: List[str]  # extracted from 不包括... segment (heuristic)


def _split_desc_for_embed(desc: str) -> str:
    """
    Embedding should not consume negation too much; remove "不包括" tail if present.
    """
    if not desc:
        return ""
    # Split on first occurrence of 不包括 / 不包含
    m = re.split(r"(不包括|不包含)", desc, maxsplit=1)
    if len(m) >= 3:
        return m[0].strip()
    return desc.strip()


def _extract_terms_after_marker(desc: str, marker: str) -> List[str]:
    """
    Heuristic term extraction from Chinese descriptions like:
      "包括A、B、C... 不包括X、Y..."
    marker: "包括" or "不包括"
    Returns a list of short phrases (Chinese or English) used for keyword routing/blocking.
    """
    if not desc or marker not in desc:
        return []
    # take substring after first marker
    idx = desc.find(marker)
    sub = desc[idx + len(marker):]

    # stop at next marker to avoid mixing
    stop_markers = ["不包括", "不包含"] if marker == "包括" else []
    for sm in stop_markers:
        j = sub.find(sm)
        if j != -1:
            sub = sub[:j]
            break

    # stop at sentence end markers
    sub = re.split(r"[。；;\n]", sub, maxsplit=1)[0]
    # remove leading punctuation
    sub = sub.strip(" ：:，,")
    if not sub:
        return []

    # split on common separators
    parts = re.split(r"[、，,/\|]+", sub)
    out = []
    for p in parts:
        t = p.strip()
        if not t:
            continue
        # remove bracketed clarifications
        t = re.sub(r"（.*?）", "", t).strip()
        t = re.sub(r"\(.*?\)", "", t).strip()
        if not t:
            continue
        # keep reasonable length
        if len(t) < 2:
            continue
        # cap
        if len(t) > 60:
            t = t[:60]
        out.append(t)
    # de-dup preserve order
    seen = set()
    dedup = []
    for t in out:
        if t.lower() in seen:
            continue
        seen.add(t.lower())
        dedup.append(t)
    return dedup


def load_classification_rules(path: str) -> List[CategoryRule]:
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"classification file not found: {p.resolve()}")

    rules: List[CategoryRule] = []
    for raw in p.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith("#"):
            continue
        if "：" in line:
            name, desc = line.split("：", 1)
        elif ":" in line:
            name, desc = line.split(":", 1)
        else:
            name, desc = line, ""
        name = name.strip()
        desc = desc.strip()
        if not name:
            continue

        desc_llm = desc
        desc_embed = _split_desc_for_embed(desc)

        include_terms = _extract_terms_after_marker(desc, "包括")
        exclude_terms = _extract_terms_after_marker(desc, "不包括") + _extract_terms_after_marker(desc, "不包含")

        rules.append(CategoryRule(
            name=name,
            desc_llm=desc_llm,
            desc_embed=desc_embed,
            include_terms=include_terms,
            exclude_terms=exclude_terms,
        ))

    if not rules:
        raise RuntimeError(f"No valid category lines parsed from {p.resolve()}")
    return rules


# ============================
# Keyword gating (include/exclude)
# ============================
def _count_term_hits(text: str, terms: List[str]) -> int:
    """
    Count how many distinct terms appear in text (case-insensitive).
    """
    if not text or not terms:
        return 0
    t = text.lower()
    hits = 0
    for term in terms:
        term_l = term.lower()
        if not term_l:
            continue
        # avoid ultra-common noise
        if len(term_l) <= 2:
            continue
        if term_l in t:
            hits += 1
    return hits


def _any_term_hit(text: str, terms: List[str]) -> bool:
    return _count_term_hits(text, terms) > 0


# ============================
# Transformer semantic routing
# ============================
def _load_embedder():
    from sentence_transformers import SentenceTransformer
    return SentenceTransformer(EMB_MODEL_NAME)


def _cosine_sim_matrix(a, b):
    """
    a: (n,d) normalized; b: (m,d) normalized
    return (n,m)
    """
    import numpy as np
    return a @ b.T


# ============================
# Gemini translation
# ============================
def google_translate_texts(texts, label: str) -> List[str]:
    if not texts:
        return []

    import requests

    out = []
    total = len(texts)
    for i, text in enumerate(texts, start=1):
        text = _normalize_cell(text)
        if not text:
            out.append("")
            continue

        last_err = None
        for attempt in range(1, GOOGLE_TRANSLATE_MAX_RETRIES + 1):
            try:
                resp = requests.get(
                    "https://translate.googleapis.com/translate_a/single",
                    params={
                        "client": "gtx",
                        "sl": "auto",
                        "tl": "zh-CN",
                        "dt": "t",
                        "q": text,
                    },
                    timeout=45,
                )
                resp.raise_for_status()
                data = resp.json()
                pieces = data[0] if isinstance(data, list) and data and isinstance(data[0], list) else []
                translated = "".join(str(p[0]) for p in pieces if isinstance(p, list) and p)
                out.append(translated.strip())
                break
            except Exception as e:
                last_err = e
                wait = min(30, 2 * attempt) + random.uniform(0, 1.5)
                print(f"[google-translate:{label}] item {i}/{total} attempt {attempt} failed: {e} (wait {wait:.1f}s)", flush=True)
                time.sleep(wait)
        else:
            raise RuntimeError(f"Google Translate failed for {label} item {i}/{total}: {last_err}")

    return out


def translate_texts(client, texts, label: str) -> List[str]:
    if not texts:
        return []

    system = (
        "You are a professional scientific translator. "
        "Translate English into Simplified Chinese. Preserve abbreviations, formulas, proper nouns, and units."
    )
    user = (
        "Return ONLY JSON in this schema: {\"translations\": [\"...\", \"...\"]}\n"
        "Array length must equal input length and order must match.\n"
        f"inputs={json.dumps(texts, ensure_ascii=False)}"
    )
    messages = [{"role": "system", "content": system}, {"role": "user", "content": user}]
    resp = _call_with_retries(client, messages, f"translate:{label}", json_object=True, max_retries=GEMINI_TRANSLATE_MAX_RETRIES)
    content = _clean_json_text(_extract_content(resp))

    try:
        data = json.loads(content)
        out = data.get("translations", [])
    except Exception:
        out = []

    if not isinstance(out, list) or len(out) != len(texts):
        raise RuntimeError(f"Translation output invalid for {label}.")
    return [str(x) for x in out]


def translate_texts_with_provider(client, texts, label: str, provider: str) -> List[str]:
    provider = (provider or "google").strip().lower()
    if provider in {"none", "skip"}:
        return ["" for _ in texts]
    if provider in {"google", "google_translate", "free"}:
        return google_translate_texts(texts, label)
    raise RuntimeError(f"Unsupported translation provider: {provider}. Use google or none.")


def enrich_translation(df: pd.DataFrame, provider: str = "google", checkpoint_path: Optional[Path] = None) -> pd.DataFrame:
    df = df.copy()
    df = ensure_base_columns(df)
    df = ensure_stable_id_column(df)
    provider = (provider or "gemini").strip().lower()

    if checkpoint_path is not None:
        df = _load_translation_checkpoint(df, checkpoint_path)

    need_title_idx = df.index[(df["title"].str.strip() != "") & (df["title_zh"].str.strip() == "")].tolist()
    need_abs_idx = df.index[(df["abstract"].str.strip() != "") & (df["abstract_zh"].str.strip() == "")].tolist()

    print(f"[plan] translation provider: {provider}", flush=True)
    print(f"[plan] titles to translate: {len(need_title_idx)}", flush=True)
    print(f"[plan] abstracts to translate: {len(need_abs_idx)} (truncate={MAX_ABSTRACT_CHARS_TO_TRANSLATE} chars)", flush=True)

    if not need_title_idx and not need_abs_idx:
        print("[plan] nothing to translate", flush=True)
        return df

    client = None

    if need_title_idx:
        start = time.time()
        total = len(need_title_idx)
        done = 0
        for batch_idx in _chunked(need_title_idx, TRANSLATE_BATCH_SIZE_TITLE):
            batch = df.loc[batch_idx, "title"].tolist()
            translated = translate_texts_with_provider(client, batch, "title", provider)
            for idx, zh in zip(batch_idx, translated):
                df.at[idx, "title_zh"] = zh
            done += len(batch_idx)
            if checkpoint_path is not None:
                _save_translation_checkpoint(df, checkpoint_path)
            elapsed = time.time() - start
            rate = done / elapsed if elapsed > 0 else 0.0
            eta = (total - done) / rate if rate > 0 else 0.0
            print(f"[translate:title] {done}/{total} | elapsed={_fmt_secs(elapsed)} | rate={rate:.2f} items/s | ETA={_fmt_secs(eta)}", flush=True)

    if need_abs_idx:
        start = time.time()
        total = len(need_abs_idx)
        done = 0
        for batch_idx in _chunked(need_abs_idx, TRANSLATE_BATCH_SIZE_ABSTRACT):
            batch = []
            for idx in batch_idx:
                x = _normalize_cell(df.at[idx, "abstract"])
                if len(x) > MAX_ABSTRACT_CHARS_TO_TRANSLATE:
                    x = x[:MAX_ABSTRACT_CHARS_TO_TRANSLATE]
                batch.append(x)
            translated = translate_texts_with_provider(client, batch, "abstract", provider)
            for idx, zh in zip(batch_idx, translated):
                df.at[idx, "abstract_zh"] = zh
            done += len(batch_idx)
            if checkpoint_path is not None:
                _save_translation_checkpoint(df, checkpoint_path)
            elapsed = time.time() - start
            rate = done / elapsed if elapsed > 0 else 0.0
            eta = (total - done) / rate if rate > 0 else 0.0
            print(f"[translate:abstract] {done}/{total} | elapsed={_fmt_secs(elapsed)} | rate={rate:.2f} items/s | ETA={_fmt_secs(eta)}", flush=True)

    return df


# ============================
# Gemini classification (by_id mapping)
# ============================
def gemini_classify_by_id(
    client,
    rules: List[CategoryRule],
    items: List[Dict[str, Any]],
    label: str,
) -> Dict[str, List[str]]:
    """
    items: list of {id, title, text, candidates(optional list[str])}
    Return: {id: [labels]}
    """
    categories_block = "\n".join([f"- {r.name}: {r.desc_llm}" for r in rules])
    category_names = [r.name for r in rules]
    allowed_set = set(category_names)

    system = (
        "You are an expert literature classifier.\n"
        "Classify each paper by the given category definitions.\n"
        "Key rules:\n"
        "1) Use abstract first; if missing, use title.\n"
        "2) Multi-label is allowed, but only when clearly justified.\n"
        "3) DO NOT classify solely by a material name (e.g., 'perovskite'). Classify by application/system.\n"
        "4) If the article does not clearly match any category definition, return empty labels [].\n"
        "5) If candidates are provided for an item, you MUST choose labels only from candidates (or []).\n"
        "6) Output MUST be valid JSON object. No extra keys.\n"
    )

    user = (
        "分类规则（类别名: 说明）:\n"
        f"{categories_block}\n\n"
        "请返回严格JSON：\n"
        "{\n"
        "  \"by_id\": {\n"
        "    \"<id>\": {\"labels\": [\"类别A\",\"类别B\"], \"confidence\": 0.0}\n"
        "  }\n"
        "}\n\n"
        "要求：\n"
        "- by_id 必须覆盖输入中所有 id。\n"
        "- labels 只能使用给定类别名；不匹配则 [].\n"
        "- confidence 为 0~1 的数字（可粗略）。\n\n"
        f"输入条目：{json.dumps(items, ensure_ascii=False)}"
    )

    messages = [{"role": "system", "content": system}, {"role": "user", "content": user}]
    resp = _call_with_retries(
        client,
        messages,
        f"classify:{label}",
        json_object=True,
        max_retries=GEMINI_CLASSIFY_MAX_RETRIES,
    )
    content = _clean_json_text(_extract_content(resp))

    try:
        data = json.loads(content)
        by_id = data.get("by_id", {})
    except Exception:
        by_id = {}

    out: Dict[str, List[str]] = {}
    if isinstance(by_id, dict):
        for _id, obj in by_id.items():
            if not isinstance(obj, dict):
                continue
            labels = obj.get("labels", [])
            if not isinstance(labels, list):
                labels = []
            clean = []
            for c in labels:
                c = _normalize_cell(c)
                if c in allowed_set:
                    clean.append(c)
            # de-dup preserve order
            out[_id] = list(dict.fromkeys(clean))

    return out


# ============================
# Hybrid classification orchestrator
# ============================
def classify_hybrid(
    df: pd.DataFrame,
    rules: List[CategoryRule],
    debug_dir: Path,
    skip_gpt: bool = False,
    checkpoint_path: Optional[Path] = None,
) -> Tuple[pd.DataFrame, List[List[str]], List[str]]:
    """
    Returns:
      df_out with categories column
      labels_list: per-row labels list aligned to df rows
      still_missing_ids: ids missing labels after 2 Gemini rounds
    """
    df = df.copy()
    df = ensure_base_columns(df)
    df = ensure_stable_id_column(df)

    # Build items
    records = df.to_dict(orient="records")
    ids = df["stable_id"].tolist()

    texts = [build_text_for_classify(r.get("title", ""), r.get("abstract", "")) for r in records]
    titles = [_normalize_cell(r.get("title", "")) for r in records]

    cat_names = [r.name for r in rules]
    name_to_rule = {r.name: r for r in rules}

    # Rule layer: exclude blocking + include strong assign
    blocked: List[set] = [set() for _ in records]
    include_hits: List[Dict[str, int]] = [dict() for _ in records]
    auto_labels: List[List[str]] = [[] for _ in records]   # from strong keyword or strong embedding
    need_model_idx: List[int] = []

    for i, text in enumerate(texts):
        t = (titles[i] + " " + text).lower()

        # compute block list
        for r in rules:
            if r.exclude_terms and _any_term_hit(t, r.exclude_terms):
                blocked[i].add(r.name)

        # include hit counts
        for r in rules:
            if r.include_terms:
                hits = _count_term_hits(t, r.include_terms)
                if hits > 0:
                    include_hits[i][r.name] = hits

        # strong include => auto label (exclude has priority)
        strong = [cn for cn, h in include_hits[i].items() if h >= KEYWORD_STRONG_HITS and cn not in blocked[i]]
        # keep order by rules list order
        strong_sorted = [cn for cn in cat_names if cn in strong]
        if strong_sorted:
            auto_labels[i] = strong_sorted
        else:
            need_model_idx.append(i)

    print(f"[gate] auto-labeled by keyword strong: {sum(1 for x in auto_labels if x)} / {len(records)}", flush=True)

    # Transformer semantic routing for remaining
    embedder = _load_embedder()

    # Build category prototype texts (embedding-friendly, no 不包括 tail)
    proto_texts = []
    for r in rules:
        # keep it positive and compact
        t = f"{r.name}。{r.desc_embed}"
        # remove "不包括" if any remained
        t = _split_desc_for_embed(t)
        proto_texts.append(t.strip())

    # Compute embeddings
    import numpy as np

    start = time.time()
    proto_emb = embedder.encode(proto_texts, normalize_embeddings=True, show_progress_bar=False)
    item_emb = embedder.encode([titles[i] + " " + texts[i] for i in range(len(records))],
                               normalize_embeddings=True, show_progress_bar=True)

    sims = _cosine_sim_matrix(np.array(item_emb), np.array(proto_emb))  # (n, C)
    print(f"[embed] computed sims in {_fmt_secs(time.time()-start)}", flush=True)

    candidates: List[List[str]] = [[] for _ in records]
    no_candidate_ids: List[str] = []

    for i in range(len(records)):
        if auto_labels[i]:
            continue

        # If category is blocked, we will exclude it from candidates and auto-assign
        sim_row = sims[i].copy()
        # Make blocked categories very low
        for j, cn in enumerate(cat_names):
            if cn in blocked[i]:
                sim_row[j] = -1.0

        # top sorted indices
        top_idx = np.argsort(-sim_row)
        top1 = top_idx[0]
        top1_score = float(sim_row[top1])

        # strong embedding => auto assign (if not blocked)
        if top1_score >= EMB_STRONG_THRESHOLD and cat_names[top1] not in blocked[i]:
            auto_labels[i] = [cat_names[top1]]
            continue

        # otherwise take topK above weak threshold
        cand = []
        for j in top_idx[: max(TOPK_CANDIDATES * 2, TOPK_CANDIDATES)]:
            sc = float(sim_row[j])
            if sc < EMB_WEAK_THRESHOLD:
                continue
            cand.append(cat_names[j])
            if len(cand) >= TOPK_CANDIDATES:
                break

        # also add weak include-hit categories (hits==1) as candidates
        weak_inc = [cn for cn, h in include_hits[i].items() if h > 0 and cn not in blocked[i]]
        for cn in cat_names:
            if cn in weak_inc and cn not in cand:
                cand.append(cn)
            if len(cand) >= TOPK_CANDIDATES:
                break

        candidates[i] = cand
        if not cand:
            no_candidate_ids.append(ids[i])

    print(f"[route] no_candidate after keyword+embed: {len(no_candidate_ids)}", flush=True)

    # Gemini final confirmation for those not auto-labeled
    final_labels: Dict[str, List[str]] = {ids[i]: auto_labels[i] for i in range(len(records)) if auto_labels[i]}
    still_missing_ids: List[str] = []
    checkpoint_labels: Dict[str, List[str]] = {}
    if checkpoint_path is not None:
        checkpoint_labels = _load_classification_checkpoint(checkpoint_path, set(cat_names))
        for _id, labs in checkpoint_labels.items():
            if _id in ids:
                final_labels[_id] = labs
        if final_labels:
            _save_classification_checkpoint(checkpoint_path, final_labels)

    if skip_gpt:
        print("[gemini] skip_gpt enabled; leaving non-auto items as []", flush=True)
        for i in range(len(records)):
            if ids[i] not in final_labels:
                final_labels[ids[i]] = []
        if checkpoint_path is not None:
            _save_classification_checkpoint(checkpoint_path, final_labels)
    else:
        client = _build_gemini_client()

        # build Gemini items for those without labels
        pending = []
        for i in range(len(records)):
            _id = ids[i]
            if _id in final_labels:
                continue

            # If no candidates, allow Gemini to choose from all, but encourage empty if unclear
            item = {
                "id": _id,
                "title": titles[i],
                "text": texts[i],
            }
            if candidates[i]:
                item["candidates"] = candidates[i]
            pending.append(item)

        print(f"[gemini] items to ask Gemini: {len(pending)}", flush=True)

        def run_batches(items_list: List[Dict[str, Any]], tag: str) -> Dict[str, List[str]]:
            out_map: Dict[str, List[str]] = {}
            start0 = time.time()
            total = len(items_list)
            done = 0
            for batch in _chunked(items_list, CLASSIFY_BATCH_SIZE):
                try:
                    resp_map = gemini_classify_by_id(client, rules, batch, label=f"{tag}:{done}")
                except Exception as e:
                    print(f"[gemini:{tag}] batch failed after retries ({len(batch)} items): {e}", flush=True)
                    resp_map = {}
                # merge
                for it in batch:
                    _id = it["id"]
                    if _id in resp_map:
                        out_map[_id] = resp_map[_id]
                        final_labels[_id] = resp_map[_id]
                if checkpoint_path is not None and final_labels:
                    _save_classification_checkpoint(checkpoint_path, final_labels)
                done = len(out_map)
                elapsed = time.time() - start0
                rate = (done / elapsed) if elapsed > 0 else 0.0
                eta = ((total - done) / rate) if rate > 0 else 0.0
                print(f"[gemini:{tag}] {done}/{total} | elapsed={_fmt_secs(elapsed)} | rate={rate:.2f} items/s | ETA={_fmt_secs(eta)}", flush=True)
            return out_map

        # first pass
        first_map = run_batches(pending, "pass1")

        # detect missing ids (漏掉)
        missing1 = [it for it in pending if it["id"] not in first_map]
        if missing1:
            # debug artifacts
            debug_dir.mkdir(parents=True, exist_ok=True)
            (debug_dir / "missing_pass1.json").write_text(json.dumps(missing1, ensure_ascii=False, indent=2), encoding="utf-8")
            print(f"[debug] dumped missing_pass1.json -> {debug_dir}", flush=True)

        # second pass only for missing
        second_map = {}
        if missing1:
            print(f"[gemini] retry missing ids (count={len(missing1)})", flush=True)
            second_map = run_batches(missing1, "pass2")

        # merge results
        for _id, labs in first_map.items():
            final_labels[_id] = labs
        for _id, labs in second_map.items():
            final_labels[_id] = labs

        # still missing after 2 passes
        still_missing_ids = [it["id"] for it in missing1 if it["id"] not in second_map]
        if still_missing_ids:
            debug_dir.mkdir(parents=True, exist_ok=True)
            (debug_dir / "still_missing_ids.json").write_text(json.dumps(still_missing_ids, ensure_ascii=False, indent=2), encoding="utf-8")
            print(f"[debug] dumped still_missing_ids.json -> {debug_dir}", flush=True)

        # fill any absent entries with []
        for it in pending:
            if it["id"] not in final_labels:
                final_labels[it["id"]] = []

    # Apply exclude blocks again to final labels (exclude > include)
    cleaned_list: List[List[str]] = []
    for i, _id in enumerate(ids):
        labs = final_labels.get(_id, [])
        labs2 = [c for c in labs if c not in blocked[i]]
        # de-dup and preserve rule order
        ordered = [cn for cn in cat_names if cn in labs2]
        cleaned_list.append(ordered)

    df["categories"] = [";".join(x) for x in cleaned_list]
    return df, cleaned_list, still_missing_ids


# ============================
# XLSX output
# ============================
def _safe_sheet_name(name: str, used: set) -> str:
    base = re.sub(r"[\\/*?:\[\]]", "_", _normalize_cell(name)) or "Sheet"
    base = base[:31]
    candidate = base
    i = 2
    while candidate in used:
        suffix = f"_{i}"
        candidate = base[: 31 - len(suffix)] + suffix
        i += 1
    used.add(candidate)
    return candidate


def write_grouped_xlsx(
    df: pd.DataFrame,
    labels_list: List[List[str]],
    ordered_categories: List[str],
    still_missing_ids: List[str],
    output_xlsx: str,
):
    used = set()
    sid = set(still_missing_ids)

    cat_to_indices = {c: [] for c in ordered_categories}
    uncategorized = []
    for i, labs in enumerate(labels_list):
        if not labs:
            uncategorized.append(i)
        for c in labs:
            if c in cat_to_indices:
                cat_to_indices[c].append(i)

    with pd.ExcelWriter(output_xlsx, engine="openpyxl") as writer:
        # ALL
        df.to_excel(writer, index=False, sheet_name=_safe_sheet_name("ALL", used))

        # category sheets
        for cat in ordered_categories:
            idxs = cat_to_indices.get(cat, [])
            if not idxs:
                continue
            sheet = _safe_sheet_name(cat, used)
            sub = df.iloc[idxs].copy()
            sub.to_excel(writer, index=False, sheet_name=sheet)

        # UNCATEGORIZED
        if uncategorized:
            sheet = _safe_sheet_name("UNCATEGORIZED", used)
            df.iloc[uncategorized].copy().to_excel(writer, index=False, sheet_name=sheet)

        # STILL_MISSING (after 2 Gemini passes)
        if sid:
            sheet = _safe_sheet_name("STILL_MISSING", used)
            df[df["stable_id"].isin(sid)].copy().to_excel(writer, index=False, sheet_name=sheet)


# ============================
# DOCX helpers (layout)
# ============================
def add_hyperlink(paragraph, url: str, text: str, color_hex="1155CC", underline=True):
    if not url:
        paragraph.add_run(text)
        return

    part = paragraph.part
    r_id = part.relate_to(
        url,
        reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color_hex)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single" if underline else "none")
    r_pr.append(u)
    new_run.append(r_pr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_doc_default_style(doc: Document, font_name="Calibri", font_size_pt=11):
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_divider_line(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("—" * 70)
    run.font.size = Pt(9)


def doi_to_url(doi: str) -> str:
    doi = _normalize_cell(doi)
    if not doi:
        return ""
    if doi.lower().startswith("http"):
        return doi
    return f"https://doi.org/{doi}"


def is_truthy_flag(s: str) -> bool:
    return _normalize_cell(s).lower() in {"true", "1", "yes", "y", "t"}


def strip_leading_abstract(text: str) -> str:
    if not text:
        return text
    s = text.lstrip()
    return re.sub(r"^(abstract)\s*[:.\-–—]*\s*", "", s, flags=re.IGNORECASE)


def abstract_to_runs(abstract: str):
    if not abstract:
        return [("(empty)", False)]

    abstract = strip_leading_abstract(abstract)
    text = abstract.replace("\r\n", "\n").replace("\r", "\n")
    lines = text.split("\n")

    chunks = []
    pending_space = False
    for line in lines:
        s = line.strip()
        if s == "":
            pending_space = True
            continue
        if pending_space and chunks:
            chunks.append((" ", False))
            pending_space = False

        if s.isdigit():
            chunks.append((s, True))
        else:
            if chunks:
                prev_text, prev_sub = chunks[-1]
                if not prev_sub and prev_text and not prev_text.endswith(" ") and prev_text[-1] not in {"-", "−", "/"}:
                    chunks.append((" ", False))
            chunks.append((s, False))

    cleaned = []
    for t, sub in chunks:
        if cleaned and t == " " and cleaned[-1][0] == " ":
            continue
        cleaned.append((t, sub))
    return cleaned


def add_abstract_with_subscripts(paragraph, abstract: str):
    for t, is_sub in abstract_to_runs(abstract):
        run = paragraph.add_run(t)
        if is_sub:
            run.font.subscript = True


def _write_record_block(doc, row, idx):
    title_en = _normalize_cell(row.get("title", ""))
    title_zh = _normalize_cell(row.get("title_zh", ""))
    link = _normalize_cell(row.get("link", ""))
    source = _normalize_cell(row.get("source", ""))
    pub_date = _normalize_cell(row.get("pub_date", "")) or _normalize_cell(row.get("published", ""))
    doi = _normalize_cell(row.get("doi", ""))
    abstract_en = _normalize_cell(row.get("abstract", ""))
    abstract_zh = _normalize_cell(row.get("abstract_zh", ""))
    abstract_source = _normalize_cell(row.get("abstract_source", ""))
    must_have_abstract = _normalize_cell(row.get("must_have_abstract", ""))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    p.add_run(f"[{idx}] ").bold = True
    main_title = title_zh or title_en or "(no title)"
    if link:
        add_hyperlink(p, link, main_title)
    else:
        p.add_run(main_title).bold = True

    if title_en and title_zh and title_en != title_zh:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(f"EN: {title_en}")
        r2.italic = True
        r2.font.size = Pt(10)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(6)
    mr = meta.add_run("Source: ")
    mr.bold = True
    meta.add_run(source or "-")
    meta.add_run("    ")
    mr = meta.add_run("Pub date: ")
    mr.bold = True
    meta.add_run(pub_date or "-")
    if doi:
        meta.add_run("    ")
        mr = meta.add_run("DOI: ")
        mr.bold = True
        add_hyperlink(meta, doi_to_url(doi), doi)

    extra_bits = []
    if abstract_source:
        extra_bits.append(f"abstract_source={abstract_source}")
    if is_truthy_flag(must_have_abstract):
        extra_bits.append("must_have_abstract=True")
    if extra_bits:
        extra = doc.add_paragraph()
        extra.paragraph_format.space_before = Pt(0)
        extra.paragraph_format.space_after = Pt(6)
        er = extra.add_run("Notes: ")
        er.bold = True
        extra.add_run("; ".join(extra_bits))

    abs_zh_p = doc.add_paragraph()
    abs_zh_p.paragraph_format.space_before = Pt(0)
    abs_zh_p.paragraph_format.space_after = Pt(4)
    abs_zh_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    abs_zh_p.paragraph_format.line_spacing = 1.15
    abs_zh_p.add_run("摘要（ZH）: ").bold = True
    add_abstract_with_subscripts(abs_zh_p, abstract_zh if abstract_zh else "(empty)")

    abs_en_p = doc.add_paragraph()
    abs_en_p.paragraph_format.space_before = Pt(0)
    abs_en_p.paragraph_format.space_after = Pt(10)
    abs_en_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    abs_en_p.paragraph_format.line_spacing = 1.15
    abs_en_p.add_run("Abstract (EN): ").bold = True
    add_abstract_with_subscripts(abs_en_p, abstract_en if abstract_en else "(empty)")


def df_to_word_bilingual_grouped(
    df: pd.DataFrame,
    labels_list: List[List[str]],
    ordered_categories: List[str],
    output_docx: str,
    report_title="Tech Tracking Digest",
):
    doc = Document()
    set_doc_default_style(doc, font_name="Calibri", font_size_pt=11)

    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(report_title)
    r.bold = True
    r.font.size = Pt(18)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(14)
    sub_p.add_run(datetime.now().strftime("%Y-%m-%d")).italic = True

    if df.empty:
        doc.add_paragraph("No records found in CSV.")
        doc.save(output_docx)
        return

    records = df.to_dict(orient="records")
    cat_to_indices = {c: [] for c in ordered_categories}
    uncategorized = []
    for i, cats in enumerate(labels_list):
        if not cats:
            uncategorized.append(i)
        for c in cats:
            if c in cat_to_indices:
                cat_to_indices[c].append(i)

    appeared_categories = [c for c in ordered_categories if cat_to_indices.get(c)]

    # write categories in order
    for ci, cat in enumerate(appeared_categories):
        header = doc.add_paragraph()
        header.paragraph_format.space_before = Pt(8)
        header.paragraph_format.space_after = Pt(6)
        hr = header.add_run(f"【{cat}】")
        hr.bold = True
        hr.font.size = Pt(14)

        indices = cat_to_indices[cat]
        for j, ridx in enumerate(indices, start=1):
            _write_record_block(doc, records[ridx], j)
            if j != len(indices):
                add_divider_line(doc)
        if ci != len(appeared_categories) - 1:
            add_divider_line(doc)

    # append uncategorized
    if uncategorized:
        add_divider_line(doc)
        header = doc.add_paragraph()
        header.paragraph_format.space_before = Pt(8)
        header.paragraph_format.space_after = Pt(6)
        hr = header.add_run("【UNCATEGORIZED】")
        hr.bold = True
        hr.font.size = Pt(14)

        for j, ridx in enumerate(uncategorized, start=1):
            _write_record_block(doc, records[ridx], j)
            if j != len(uncategorized):
                add_divider_line(doc)

    doc.save(output_docx)


# ============================
# main
# ============================
def main():
    global GOOGLE_AI_MODEL
    parser = argparse.ArgumentParser()
    parser.add_argument("-i", "--input", default="", help="Input weekly CSV path (default: latest in output/weekly)")
    parser.add_argument("-c", "--classification", default=DEFAULT_CLASSIFICATION_FILE, help="Classification rules file")
    parser.add_argument("--report-title", default="Tech Tracking Digest")
    parser.add_argument("--skip-translate", action="store_true", help="Skip translation")
    parser.add_argument(
        "--translation-provider",
        default=TRANSLATE_PROVIDER,
        choices=["google", "none"],
        help="Translation route: google or none. Classification still uses Gemini.",
    )
    parser.add_argument(
        "--classification-model",
        default=GOOGLE_AI_MODEL,
        choices=["gemini-2.5-flash", "gemini-2.5-flash-lite"],
        help="Gemini model for classification only.",
    )
    parser.add_argument("--output-suffix", default="_translated", help="Suffix inserted before .xlsx/.docx")
    parser.add_argument("--save-translated-csv", action="store_true", help="Save translated rows before classification")
    parser.add_argument("--checkpoint-dir", default="output/checkpoints", help="Directory for resumable translation/classification checkpoints")
    parser.add_argument("--no-resume", action="store_true", help="Ignore existing checkpoints for this run")
    parser.add_argument("--skip-gpt", action="store_true", help="Skip Gemini classification (only keyword+embedding)")
    parser.add_argument("--debug-dir", default="output/debug", help="Debug folder for failure artifacts")
    args = parser.parse_args()

    GOOGLE_AI_MODEL = args.classification_model.strip()
    if not GOOGLE_AI_MODEL:
        raise RuntimeError("GOOGLE_AI_MODEL is empty. Set GOOGLE_AI_MODEL or GEMINI_MODEL env.")

    if args.input:
        csv_path = Path(args.input)
    else:
        csv_path = pick_latest_weekly_csv(folder="output/weekly")

    print(f"[io] Picked CSV: {csv_path}", flush=True)
    df = pd.read_csv(csv_path, encoding="utf-8-sig", keep_default_na=False)
    df = ensure_base_columns(df)
    print(f"[io] Loaded rows: {len(df)}", flush=True)

    rules = load_classification_rules(args.classification)
    ordered_categories = [r.name for r in rules]
    print(f"[classify] loaded categories ({len(ordered_categories)}): {ordered_categories}", flush=True)

    output_suffix = args.output_suffix if args.output_suffix.startswith("_") else "_" + args.output_suffix
    checkpoint_base = Path(args.checkpoint_dir) / f"{csv_path.stem}{output_suffix}"
    translation_checkpoint = checkpoint_base.with_name(checkpoint_base.name + "_translation.csv")
    classification_checkpoint = checkpoint_base.with_name(checkpoint_base.name + "_classification.json")
    if args.no_resume:
        translation_checkpoint = None
        classification_checkpoint = None
        print("[checkpoint] resume disabled by --no-resume", flush=True)
    else:
        print(f"[checkpoint] translation: {translation_checkpoint}", flush=True)
        print(f"[checkpoint] classification: {classification_checkpoint}", flush=True)

    translation_provider = "none" if args.skip_translate else args.translation_provider
    if translation_provider != "none":
        # Translation uses Google Translate only. Classification still uses English title/abstract below.
        df = enrich_translation(df, provider=translation_provider, checkpoint_path=translation_checkpoint)
    else:
        if args.skip_translate:
            print("[plan] skip translation by --skip-translate", flush=True)
        if translation_provider == "none":
            print("[plan] skip translation by --translation-provider=none", flush=True)
        if args.skip_gpt:
            print("[plan] skip Gemini classification by --skip-gpt", flush=True)

    debug_dir = Path(args.debug_dir)

    if args.save_translated_csv:
        translated_csv = csv_path.with_name(csv_path.stem + output_suffix + ".csv")
        df.to_csv(translated_csv, index=False, encoding="utf-8-sig")
        print(f"[io] Wrote translated CSV checkpoint: {translated_csv}", flush=True)

    df2, labels_list, still_missing_ids = classify_hybrid(
        df=df,
        rules=rules,
        debug_dir=debug_dir,
        skip_gpt=args.skip_gpt,
        checkpoint_path=classification_checkpoint,
    )

    output_xlsx = csv_path.with_name(csv_path.stem + output_suffix + ".xlsx")
    write_grouped_xlsx(df2, labels_list, ordered_categories, still_missing_ids, str(output_xlsx))
    print(f"[io] Wrote XLSX: {output_xlsx}", flush=True)

    output_docx = output_xlsx.with_suffix(".docx")
    df_to_word_bilingual_grouped(df2, labels_list, ordered_categories, str(output_docx), report_title=args.report_title)
    print(f"[io] Wrote DOCX: {output_docx}", flush=True)


if __name__ == "__main__":
    main()
